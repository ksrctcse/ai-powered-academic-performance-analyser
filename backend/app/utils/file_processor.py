"""
File processor utility for extracting text from multiple file formats.
Supports PDF, DOCX, and CSV file formats.
"""

import io
import pandas as pd
from pathlib import Path
from typing import Tuple
from ..core.logger import get_logger

logger = get_logger(__name__)

# File size limits (in bytes)
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'csv', 'txt'}


class FileProcessingError(Exception):
    """Custom exception for file processing errors"""
    pass


def validate_file(filename: str, file_size: int) -> None:
    """
    Validate file before processing.
    
    Args:
        filename: Name of the file
        file_size: Size of the file in bytes
        
    Raises:
        FileProcessingError: If file is invalid
    """
    if not filename:
        raise FileProcessingError("Filename cannot be empty")
    
    # Check file extension
    file_ext = Path(filename).suffix.lower().lstrip('.')
    if file_ext not in ALLOWED_EXTENSIONS:
        raise FileProcessingError(
            f"File type '.{file_ext}' not supported. "
            f"Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    # Check file size
    if file_size > MAX_FILE_SIZE:
        raise FileProcessingError(
            f"File size ({file_size / 1024 / 1024:.2f} MB) exceeds maximum allowed size "
            f"({MAX_FILE_SIZE / 1024 / 1024:.2f} MB)"
        )
    
    if file_size == 0:
        raise FileProcessingError("File is empty")


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file_content: Binary content of the PDF file
        
    Returns:
        Extracted text from PDF
        
    Raises:
        FileProcessingError: If PDF processing fails
    """
    try:
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        
        if not pdf_reader.pages:
            raise FileProcessingError("PDF file is empty or corrupted")
        
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    # Add page text without adding page header markers
                    # This keeps the text clean for analysis
                    if text:  # Add separator between pages if already have content
                        text += "\n\n"
                    text += page_text
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                # Continue with next page
                continue
        
        if not text.strip():
            raise FileProcessingError("No text could be extracted from the PDF")
        
        return text.strip()
        
    except FileProcessingError:
        raise
    except ImportError:
        raise FileProcessingError(
            "PyPDF2 library is not installed. Please install it using: pip install PyPDF2"
        )
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}", exc_info=True)
        raise FileProcessingError(f"Failed to process PDF file: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_content: Binary content of the DOCX file
        
    Returns:
        Extracted text from DOCX
        
    Raises:
        FileProcessingError: If DOCX processing fails
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_content))
        
        if not doc.paragraphs:
            raise FileProcessingError("DOCX file is empty or contains no text")
        
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        if not text.strip():
            raise FileProcessingError("No text could be extracted from the DOCX file")
        
        return text.strip()
        
    except FileProcessingError:
        raise
    except ImportError:
        raise FileProcessingError(
            "python-docx library is not installed. Please install it using: pip install python-docx"
        )
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}", exc_info=True)
        raise FileProcessingError(f"Failed to process DOCX file: {str(e)}")


def extract_text_from_csv(file_content: bytes) -> str:
    """
    Extract text from CSV file.
    
    Args:
        file_content: Binary content of the CSV file
        
    Returns:
        Extracted text from CSV (formatted for readability)
        
    Raises:
        FileProcessingError: If CSV processing fails
    """
    try:
        # Try to read CSV with different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        df = None
        
        for encoding in encodings:
            try:
                df = pd.read_csv(io.BytesIO(file_content), encoding=encoding)
                break
            except (UnicodeDecodeError, Exception):
                continue
        
        if df is None or df.empty:
            raise FileProcessingError("CSV file is empty or cannot be read")
        
        # Convert DataFrame to formatted text
        text = "CSV Content:\n\n"
        text += df.to_string(index=False)
        
        return text
        
    except FileProcessingError:
        raise
    except ImportError:
        raise FileProcessingError(
            "pandas library is not installed. Please install it using: pip install pandas"
        )
    except Exception as e:
        logger.error(f"Error extracting text from CSV: {str(e)}", exc_info=True)
        raise FileProcessingError(f"Failed to process CSV file: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from TXT file.
    
    Args:
        file_content: Binary content of the TXT file
        
    Returns:
        Extracted text from TXT
        
    Raises:
        FileProcessingError: If TXT processing fails
    """
    try:
        # Try multiple encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                if text.strip():
                    return text.strip()
            except (UnicodeDecodeError, Exception):
                continue
        
        raise FileProcessingError("Could not decode text file with any supported encoding")
        
    except FileProcessingError:
        raise
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}", exc_info=True)
        raise FileProcessingError(f"Failed to process TXT file: {str(e)}")


def process_file(filename: str, file_content: bytes) -> Tuple[str, str]:
    """
    Process uploaded file and extract text content.
    
    Args:
        filename: Name of the uploaded file
        file_content: Binary content of the file
        
    Returns:
        Tuple of (extracted_text, file_type)
        
    Raises:
        FileProcessingError: If file processing fails
    """
    try:
        # Validate file
        validate_file(filename, len(file_content))
        
        # Get file extension
        file_ext = Path(filename).suffix.lower().lstrip('.')
        
        logger.info(f"Processing file: {filename} (size: {len(file_content)} bytes, type: {file_ext})")
        
        # Extract text based on file type
        if file_ext == 'pdf':
            text = extract_text_from_pdf(file_content)
        elif file_ext == 'docx':
            text = extract_text_from_docx(file_content)
        elif file_ext == 'csv':
            text = extract_text_from_csv(file_content)
        elif file_ext == 'txt':
            text = extract_text_from_txt(file_content)
        else:
            raise FileProcessingError(f"Unsupported file type: {file_ext}")
        
        logger.info(f"Successfully processed file: {filename} (extracted {len(text)} characters)")
        
        return text, file_ext
        
    except FileProcessingError:
        raise
    except Exception as e:
        logger.error(f"Unexpected error processing file {filename}: {str(e)}", exc_info=True)
        raise FileProcessingError(f"Failed to process file: {str(e)}")


def extract_course_name_from_text(text: str, filename: str) -> str:
    """
    Extract or infer course name from text or filename.
    
    Args:
        text: Extracted text content
        filename: Original filename
        
    Returns:
        Course name (inferred or extracted)
    """
    # Try to get from filename first
    if filename:
        # Remove extension and use as course name
        course_name = Path(filename).stem
        if course_name and len(course_name) > 2:
            return course_name
    
    # Try to extract from first line of text
    lines = text.split('\n')
    for line in lines[:5]:
        stripped = line.strip()
        if stripped and len(stripped) > 5 and not stripped.isupper():
            return stripped[:100]  # First 100 chars
    
    # Default
    return "Syllabus"
